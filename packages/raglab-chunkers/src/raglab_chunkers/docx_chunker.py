"""
DOCXChunker — heading-aware DOCX chunking with boundary backtracking.

Strategy (R2 FRS spec):
    1. Parse DOCX using python-docx, extracting paragraphs with style info.
    2. Identify heading paragraphs (Heading 1/2/3/etc.) as structural boundaries.
    3. Group body paragraphs under their nearest preceding heading (structural unit).
    4. Run `split_into_windows()` within each structural unit.
    5. If `include_heading_in_chunk=True`, prepend the heading text to every
       chunk produced from that section — downstream retrieval retains context.

Parameters:
    chunk_size             : int   = 500
    chunk_overlap          : int   = 50
    boundary_enforcement   : bool  = True
    boundary_chars         : list  = ['.','!','?']
    tokenizer              : str   = "tiktoken" | "word_count"
    min_chunk_size         : int   = 50
    preserve_headings      : bool  = True  — split at heading boundaries
    include_heading_in_chunk: bool = True  — prepend heading text to each chunk

Reuse rule: token+boundary splitting via `_boundary.split_into_windows()` only.
"""

from __future__ import annotations

import uuid
from typing import Any

from raglab_common.exceptions import ChunkerError
from raglab_common.models import ChunkModel

from raglab_chunkers._boundary import count_tokens, split_into_windows
from raglab_chunkers.base import BaseChunker

# Heading style name prefixes recognised by python-docx
_HEADING_PREFIXES = ("Heading", "heading", "Title", "title")


def _is_heading(paragraph: Any) -> bool:
    """Return True if the paragraph has a heading style."""
    style_name = paragraph.style.name if paragraph.style else ""
    return any(style_name.startswith(p) for p in _HEADING_PREFIXES)


def _heading_level(paragraph: Any) -> int:
    """Extract numeric heading level (1-9), or 0 for non-headings."""
    style_name = paragraph.style.name if paragraph.style else ""
    for prefix in _HEADING_PREFIXES:
        if style_name.startswith(prefix):
            rest = style_name[len(prefix):].strip()
            if rest.isdigit():
                return int(rest)
            return 1  # "Heading" or "Title" with no number → treat as H1
    return 0


class DOCXChunker(BaseChunker):
    """
    Heading-aware DOCX chunker.

    Splits at heading boundaries, then applies token+boundary backtracking
    within each section. Requires python-docx. Activates in R2.
    """

    chunker_type: str = "docx"

    _DEFAULT_CHUNK_SIZE: int = 500
    _DEFAULT_CHUNK_OVERLAP: int = 50
    _DEFAULT_BOUNDARY_ENFORCEMENT: bool = True
    _DEFAULT_BOUNDARY_CHARS: list[str] = [".", "!", "?"]
    _DEFAULT_TOKENIZER: str = "tiktoken"
    _DEFAULT_MIN_CHUNK_SIZE: int = 50
    _DEFAULT_PRESERVE_HEADINGS: bool = True
    _DEFAULT_INCLUDE_HEADING_IN_CHUNK: bool = True

    def __init__(self, config: dict[str, Any] | None = None) -> None:
        super().__init__(config)
        cfg = config or {}
        self.chunk_size: int = int(cfg.get("chunk_size", self._DEFAULT_CHUNK_SIZE))
        self.chunk_overlap: int = int(cfg.get("chunk_overlap", self._DEFAULT_CHUNK_OVERLAP))
        self.boundary_enforcement: bool = bool(
            cfg.get("boundary_enforcement", self._DEFAULT_BOUNDARY_ENFORCEMENT)
        )
        self.boundary_chars: frozenset[str] = frozenset(
            cfg.get("boundary_chars", self._DEFAULT_BOUNDARY_CHARS)
        )
        self.tokenizer: str = cfg.get("tokenizer", self._DEFAULT_TOKENIZER)
        self.min_chunk_size: int = int(cfg.get("min_chunk_size", self._DEFAULT_MIN_CHUNK_SIZE))
        self.preserve_headings: bool = bool(
            cfg.get("preserve_headings", self._DEFAULT_PRESERVE_HEADINGS)
        )
        self.include_heading_in_chunk: bool = bool(
            cfg.get("include_heading_in_chunk", self._DEFAULT_INCLUDE_HEADING_IN_CHUNK)
        )

        if self.chunk_size < 1:
            raise ValueError(f"chunk_size must be >= 1, got {self.chunk_size}")
        if self.chunk_overlap < 0:
            raise ValueError(f"chunk_overlap must be >= 0, got {self.chunk_overlap}")
        if self.chunk_overlap >= self.chunk_size:
            raise ValueError(
                f"chunk_overlap ({self.chunk_overlap}) must be < chunk_size ({self.chunk_size})"
            )
        if self.tokenizer not in ("tiktoken", "word_count"):
            raise ValueError(f"tokenizer must be 'tiktoken' or 'word_count', got {self.tokenizer!r}")

    def chunk_docx_bytes(
        self, docx_bytes: bytes, doc_id: str, metadata: dict[str, Any] | None = None
    ) -> list[ChunkModel]:
        """Chunk from raw DOCX bytes (BytesIO-compatible)."""
        import io
        try:
            import docx as python_docx
        except ImportError as exc:
            raise ChunkerError("python-docx not installed. Run: pip install python-docx") from exc

        try:
            doc = python_docx.Document(io.BytesIO(docx_bytes))
        except Exception as exc:
            raise ChunkerError(f"Failed to open DOCX: {exc}") from exc

        return self._process_document(doc, doc_id, metadata or {})

    def _chunk(self, text: str, doc_id: str, metadata: dict[str, Any]) -> list[ChunkModel]:
        """
        Chunk from plain text (no heading structure available).

        Falls back to single-unit split_into_windows — same as TextChunker.
        Heading structure is only available when using chunk_docx_bytes().
        """
        raw_chunks = split_into_windows(
            text=text,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            boundary_enforcement=self.boundary_enforcement,
            boundary_chars=self.boundary_chars,
            tokenizer=self.tokenizer,
            min_chunk_size=self.min_chunk_size,
        )
        return [
            ChunkModel(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                text=chunk_text,
                chunk_index=i,
                token_count=count_tokens(chunk_text, mode=self.tokenizer),
                metadata={**metadata, "chunker": self.chunker_type, "tokenizer": self.tokenizer},
            )
            for i, chunk_text in enumerate(raw_chunks)
        ]

    def _process_document(
        self, doc: Any, doc_id: str, metadata: dict[str, Any]
    ) -> list[ChunkModel]:
        """Extract heading-grouped sections and chunk each with split_into_windows."""
        if not self.preserve_headings:
            # No heading awareness — treat entire document as one stream
            full_text = "\n\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            return self._chunk(full_text, doc_id, metadata)

        # Group paragraphs by section under their heading
        sections: list[dict[str, Any]] = []
        current_heading: str | None = None
        current_heading_level: int = 0
        current_body: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if _is_heading(para):
                # Flush current section
                if current_body:
                    sections.append({
                        "heading": current_heading,
                        "heading_level": current_heading_level,
                        "body": " ".join(current_body),
                    })
                    current_body = []
                current_heading = text
                current_heading_level = _heading_level(para)
            else:
                current_body.append(text)

        # Flush final section
        if current_body:
            sections.append({
                "heading": current_heading,
                "heading_level": current_heading_level,
                "body": " ".join(current_body),
            })

        # Also handle case where document has no headings
        if not sections:
            full_text = "\n\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            return self._chunk(full_text, doc_id, metadata)

        # Chunk each section with split_into_windows
        chunks: list[ChunkModel] = []
        chunk_index = 0

        for section in sections:
            body = section["body"]
            heading = section["heading"]

            # Prepend heading to section text if configured
            section_text = f"{heading}\n{body}" if (heading and self.include_heading_in_chunk) else body
            if not section_text.strip():
                continue

            raw_chunks = split_into_windows(
                text=section_text,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                boundary_enforcement=self.boundary_enforcement,
                boundary_chars=self.boundary_chars,
                tokenizer=self.tokenizer,
                min_chunk_size=self.min_chunk_size,
            )

            for chunk_text in raw_chunks:
                chunk_meta = {
                    **metadata,
                    "chunker": self.chunker_type,
                    "tokenizer": self.tokenizer,
                    "boundary_enforcement": self.boundary_enforcement,
                }
                if heading:
                    chunk_meta["heading"] = heading
                    chunk_meta["heading_level"] = section["heading_level"]

                chunks.append(ChunkModel(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    text=chunk_text,
                    chunk_index=chunk_index,
                    token_count=count_tokens(chunk_text, mode=self.tokenizer),
                    metadata=chunk_meta,
                ))
                chunk_index += 1

        return chunks

    @classmethod
    def config_schema(cls) -> dict[str, Any]:
        return {
            "chunk_size": {
                "type": "int", "default": cls._DEFAULT_CHUNK_SIZE,
                "min": 50, "max": 4000,
                "description": "Target token count per chunk.",
            },
            "chunk_overlap": {
                "type": "int", "default": cls._DEFAULT_CHUNK_OVERLAP,
                "min": 0, "max": 500,
                "description": "Overlap between consecutive chunks.",
            },
            "boundary_enforcement": {
                "type": "bool", "default": cls._DEFAULT_BOUNDARY_ENFORCEMENT,
                "description": "Backtrack to nearest sentence boundary.",
            },
            "boundary_chars": {
                "type": "list", "default": cls._DEFAULT_BOUNDARY_CHARS,
                "description": "Characters treated as sentence boundaries.",
            },
            "tokenizer": {
                "type": "str", "default": cls._DEFAULT_TOKENIZER,
                "options": ["tiktoken", "word_count"],
                "description": "Token counting mode.",
            },
            "min_chunk_size": {
                "type": "int", "default": cls._DEFAULT_MIN_CHUNK_SIZE,
                "min": 1, "max": 200,
                "description": "Minimum tokens per chunk.",
            },
            "preserve_headings": {
                "type": "bool", "default": cls._DEFAULT_PRESERVE_HEADINGS,
                "description": "Split at DOCX heading boundaries before token chunking.",
            },
            "include_heading_in_chunk": {
                "type": "bool", "default": cls._DEFAULT_INCLUDE_HEADING_IN_CHUNK,
                "description": "Prepend heading text to every chunk in that section.",
            },
        }
